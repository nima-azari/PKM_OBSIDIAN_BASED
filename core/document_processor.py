"""
Document Processing Pipeline

Handles various document formats (txt, md, PDF) and adds them to the vault.
"""

import os
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
import hashlib
import re

import pypdf
from bs4 import BeautifulSoup
import html2text
from youtube_transcript_api import YouTubeTranscriptApi
import core.obsidian_api as obs_api

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocumentProcessor:
    """Process and ingest documents into vault"""
    
    def __init__(self, vault_path: str = "10-Projects/Cloud-vs-KG-Data-Centric/Literature"):
        self.vault_path = vault_path
    
    def process_file(self, file_path: str, tags: Optional[List[str]] = None) -> str:
        """Process file and add to vault"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Extract content based on file type
        if file_path.suffix.lower() == '.pdf':
            content = self._extract_pdf(file_path)
        elif file_path.suffix.lower() in ['.txt', '.md']:
            content = self._extract_text(file_path)
        elif file_path.suffix.lower() in ['.html', '.htm']:
            content = self._extract_html(file_path)
        elif file_path.suffix.lower() == '.docx':
            content = self._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")
        
        # Create note in vault
        note_path = self._create_note(
            file_path.stem,
            content,
            source_file=str(file_path),
            tags=tags
        )
        
        return note_path
    
    def _extract_pdf(self, pdf_path: Path) -> str:
        """Extract text from PDF"""
        reader = pypdf.PdfReader(pdf_path)
        
        text_parts = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text.strip():
                text_parts.append(f"## Page {i+1}\n\n{text}")
        
        return "\n\n".join(text_parts)
    
    def _extract_text(self, text_path: Path) -> str:
        """Extract text from txt/md file"""
        with open(text_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _extract_html(self, html_path: Path) -> str:
        """Extract text from HTML file"""
        with open(html_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # Parse HTML
        soup = BeautifulSoup(html_content, 'lxml')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Convert to markdown
        h = html2text.HTML2Text()
        h.ignore_links = False
        h.body_width = 0  # Don't wrap lines
        markdown_content = h.handle(str(soup))
        
        return markdown_content.strip()
    
    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file by converting to TXT."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        try:
            doc = DocxDocument(str(file_path))
            
            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        table_text.append(row_text)
            
            # Combine all text
            all_text = '\n\n'.join(paragraphs)
            if table_text:
                all_text += '\n\n## Tables\n\n' + '\n'.join(table_text)
            
            return all_text.strip()
        except Exception as e:
            raise ValueError(f"Could not extract DOCX content: {e}")
    
    def process_youtube_url(self, url: str, tags: Optional[List[str]] = None, save_to_sources: bool = True, convert_to_article: bool = False) -> str:
        """Extract transcript from YouTube video and save to sources directory
        
        Args:
            url: YouTube video URL
            tags: Optional tags for the document
            save_to_sources: If True, saves to data/sources/, else uses vault_path
            convert_to_article: If True, converts timestamped transcript to continuous article format using AI
        """
        # Extract video ID from URL
        video_id = self._extract_youtube_id(url)
        if not video_id:
            raise ValueError(f"Invalid YouTube URL: {url}")
        
        try:
            # Get transcript using instance method
            api = YouTubeTranscriptApi()
            transcript_obj = api.fetch(video_id)
            
            # Format transcript - transcript_obj is a FetchedTranscript with start, text, duration attributes
            transcript_text = []
            raw_text_parts = []
            
            for entry in transcript_obj:
                timestamp = self._format_timestamp(entry.start)
                text = entry.text
                transcript_text.append(f"**[{timestamp}]** {text}")
                raw_text_parts.append(text)
            
            # Choose format based on convert_to_article flag
            if convert_to_article:
                raw_transcript = " ".join(raw_text_parts)
                content = self._convert_transcript_to_article(raw_transcript, url)
            else:
                content = "\n\n".join(transcript_text)
            
            # Get video title (use video_id for now)
            title = f"YouTube-{video_id}"
            
            if save_to_sources:
                # Save directly to sources directory as markdown
                return self._save_to_sources(
                    title=title,
                    content=content,
                    url=url,
                    video_id=video_id,
                    tags=tags if tags else ['youtube', 'video', 'transcript']
                )
            else:
                # Use vault path (legacy behavior)
                metadata = {
                    'source_url': url,
                    'video_id': video_id,
                    'type': 'youtube-transcript'
                }
                
                note_path = self.add_text_note(
                    title=title,
                    content=content,
                    tags=tags if tags else ['youtube', 'video', 'transcript'],
                    metadata=metadata
                )
                
                return note_path
            
        except Exception as e:
            raise ValueError(f"Could not extract YouTube transcript: {e}")
    
    def _convert_transcript_to_article(self, raw_transcript: str, url: str) -> str:
        """Convert raw transcript to article format using AI"""
        try:
            from openai import OpenAI
            client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
            
            system_prompt = """You are a professional content writer who converts video transcripts into well-structured articles.

Your task is to:
1. Clean up the raw transcript (fix grammar, remove filler words, fix run-on sentences)
2. Organize content into logical sections with clear headings
3. Add structure: Introduction, main points, conclusion
4. Preserve all important information and key points
5. Make it readable and professional
6. Use markdown formatting with proper headings (##, ###)

Do NOT add information that wasn't in the original transcript.
Output ONLY the article content in markdown format."""

            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"Convert this YouTube transcript into a well-structured article:\n\n{raw_transcript}"}
                ],
                temperature=0.7,
                max_tokens=4000
            )
            
            article_content = response.choices[0].message.content.strip()
            
            # Add source reference
            article_content += f"\n\n---\n\n**Source:** [YouTube Video]({url})"
            
            return article_content
            
        except Exception as e:
            print(f"Warning: Could not convert to article format: {e}")
            print("Falling back to raw transcript format")
            return raw_transcript
    
    def _save_to_sources(self, title: str, content: str, url: str, video_id: str, tags: List[str]) -> str:
        """Save content directly to sources directory as markdown file"""
        # Build frontmatter
        frontmatter_lines = ["---"]
        frontmatter_lines.append(f"title: {title}")
        frontmatter_lines.append(f"source_url: {url}")
        frontmatter_lines.append(f"video_id: {video_id}")
        frontmatter_lines.append(f"type: youtube-transcript")
        frontmatter_lines.append(f"created: {datetime.now().isoformat()}")
        
        if tags:
            tags_str = ", ".join(tags)
            frontmatter_lines.append(f"tags: [{tags_str}]")
        
        frontmatter_lines.append("---")
        
        # Combine frontmatter and content
        full_content = "\n".join(frontmatter_lines) + "\n\n# " + title + "\n\n" + content
        
        # Save to sources directory
        safe_title = self._sanitize_filename(title)
        sources_dir = Path("data/sources")
        sources_dir.mkdir(parents=True, exist_ok=True)
        
        file_path = sources_dir / f"{safe_title}.md"
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(full_content)
        
        # Note: Caller will print the path
        return str(file_path)
    
    def _extract_youtube_id(self, url: str) -> Optional[str]:
        """Extract video ID from YouTube URL"""
        # Patterns for different YouTube URL formats
        patterns = [
            r'(?:https?://)?(?:www\.)?youtube\.com/watch\?v=([a-zA-Z0-9_-]{11})',
            r'(?:https?://)?(?:www\.)?youtu\.be/([a-zA-Z0-9_-]{11})',
            r'(?:https?://)?(?:www\.)?youtube\.com/embed/([a-zA-Z0-9_-]{11})',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, url)
            if match:
                return match.group(1)
        
        return None
    
    def _format_timestamp(self, seconds: float) -> str:
        """Format seconds to HH:MM:SS or MM:SS"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        
        if hours > 0:
            return f"{hours:02d}:{minutes:02d}:{secs:02d}"
        else:
            return f"{minutes:02d}:{secs:02d}"
    
    def _create_note(
        self,
        title: str,
        content: str,
        source_file: Optional[str] = None,
        tags: Optional[List[str]] = None
    ) -> str:
        """Create note in vault with frontmatter"""
        
        # Sanitize title
        safe_title = self._sanitize_filename(title)
        
        # Build frontmatter
        frontmatter_lines = ["---"]
        frontmatter_lines.append(f"title: {title}")
        frontmatter_lines.append(f"type: literature")
        frontmatter_lines.append(f"created: {datetime.now().isoformat()}")
        
        if source_file:
            frontmatter_lines.append(f"source_file: {source_file}")
        
        if tags:
            tags_str = ", ".join(tags)
            frontmatter_lines.append(f"tags: [{tags_str}]")
        
        frontmatter_lines.append("---")
        
        # Combine frontmatter and content
        full_content = "\n".join(frontmatter_lines) + "\n\n" + content
        
        # Create file in vault
        note_path = f"{self.vault_path}/{safe_title}.md"
        
        try:
            # Try to create new file
            obs_api.put_file(note_path, full_content)
            print(f"Created note: {note_path}")
        except Exception as e:
            print(f"Warning: Could not create note via API: {e}")
            print("Note content prepared but not uploaded.")
        
        return note_path
    
    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename for vault"""
        # Remove invalid characters
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            filename = filename.replace(char, '')
        
        # Limit length
        if len(filename) > 100:
            filename = filename[:100]
        
        return filename.strip()
    
    def process_directory(
        self,
        directory_path: str,
        tags: Optional[List[str]] = None,
        recursive: bool = False
    ) -> List[str]:
        """Process all supported files in directory"""
        directory = Path(directory_path)
        
        if not directory.exists() or not directory.is_dir():
            raise ValueError(f"Invalid directory: {directory_path}")
        
        supported_extensions = ['.txt', '.md', '.pdf', '.html', '.htm']
        created_notes = []
        
        # Find files
        if recursive:
            files = []
            for ext in supported_extensions:
                files.extend(directory.rglob(f'*{ext}'))
        else:
            files = []
            for ext in supported_extensions:
                files.extend(directory.glob(f'*{ext}'))
        
        # Process each file
        for file_path in files:
            try:
                note_path = self.process_file(file_path, tags=tags)
                created_notes.append(note_path)
            except Exception as e:
                print(f"Error processing {file_path}: {e}")
        
        return created_notes
    
    def add_text_note(
        self,
        title: str,
        content: str,
        tags: Optional[List[str]] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> str:
        """Add text content directly as note"""
        
        # Build frontmatter
        frontmatter_lines = ["---"]
        frontmatter_lines.append(f"title: {title}")
        frontmatter_lines.append(f"type: note")
        frontmatter_lines.append(f"created: {datetime.now().isoformat()}")
        
        if tags:
            tags_str = ", ".join(tags)
            frontmatter_lines.append(f"tags: [{tags_str}]")
        
        if metadata:
            for key, value in metadata.items():
                frontmatter_lines.append(f"{key}: {value}")
        
        frontmatter_lines.append("---")
        
        # Combine
        full_content = "\n".join(frontmatter_lines) + "\n\n" + content
        
        # Create in vault
        safe_title = self._sanitize_filename(title)
        note_path = f"{self.vault_path}/{safe_title}.md"
        
        try:
            obs_api.put_file(note_path, full_content)
            print(f"Created note: {note_path}")
        except Exception as e:
            print(f"Warning: Could not create note: {e}")
        
        return note_path
    
    def get_document_hash(self, file_path: str) -> str:
        """Get hash of document for deduplication"""
        hasher = hashlib.md5()
        
        with open(file_path, 'rb') as f:
            buf = f.read()
            hasher.update(buf)
        
        return hasher.hexdigest()


if __name__ == "__main__":
    # Demo
    processor = DocumentProcessor()
    
    print("Document Processor initialized")
    print(f"Target vault path: {processor.vault_path}")
